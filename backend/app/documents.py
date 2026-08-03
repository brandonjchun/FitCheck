"""
Raw text extraction from uploaded resume files.

Pure functions: bytes in, text out. No database, no network, no framework.
"""

from collections.abc import Iterator
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

import docx
import pdfplumber
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from pdfplumber.utils.exceptions import PdfminerException

# Cells of one row are joined with this when the whole row fits on one line.
# A resume skills table -- "Languages | Python, Go" -- only means anything if
# the label stays attached to its values, and a bare newline between cells
# breaks exactly that association.
CELL_SEPARATOR = " | "


class DocumentError(Exception):
    """Base for every failure caused by the *input document*.

    The point of a common base is that the caller can write one `except
    DocumentError` and map the whole category to 400, without importing
    pdfplumber or zipfile to know what a bad upload looks like. Anything that
    is not a DocumentError is a bug in our code and should surface as a 500.

    This is the same error-classification idea the queue uses in M6, where
    failures split into "retry this" and "never retry this" -- classify by
    what the caller should *do*, not by where the error happened.
    """


class UnsupportedFileTypeError(DocumentError):
    """Raised when the uploaded file is not a format we can extract text from."""


class CorruptDocumentError(DocumentError):
    """Raised when the format is supported but the bytes cannot be parsed.

    Truncated uploads, garbage renamed to .pdf, a .docx that is not really a
    Word file. The client sent something unreadable -- that is their problem
    to fix, not a server fault.
    """


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from an uploaded resume.

    Args:
        file_bytes: The full contents of the uploaded file.
        filename: The original filename, e.g. "brandon_resume.pdf".

    Returns:
        The document's text content as a single string.

    Raises:
        UnsupportedFileTypeError: If the file is not a PDF or DOCX.
    """
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return _extract_pdf(file_bytes)
    if extension == ".docx":
        return _extract_docx(file_bytes)

    raise UnsupportedFileTypeError(
        f"Cannot extract text from '{extension or 'no extension'}': {filename}"
    )


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF bytes.

    Raises:
        CorruptDocumentError: If the bytes cannot be parsed as a PDF.
    """
    # pdfplumber needs a file-like object -- something with .read/.seek/.tell.
    # BytesIO wraps the bytes we already have in exactly that interface,
    # backed by memory instead of disk. Nothing touches the filesystem.
    buffer = BytesIO(file_bytes)

    # Only open() is inside the try. pdfminer validates the header and the
    # cross-reference table here, which is where every malformed PDF fails.
    # Keeping the loop outside means a bug in our own extraction code
    # surfaces as itself instead of being mislabelled a corrupt upload.
    try:
        pdf = pdfplumber.open(buffer)
    except PdfminerException as exc:
        raise CorruptDocumentError(
            "This file could not be read as a PDF. It may be corrupt, "
            "incomplete, or not actually a PDF."
        ) from exc

    page_texts: list[str] = []

    with pdf:
        for page in pdf.pages:
            text = page.extract_text()
            # Pages with no text layer (scanned images) return None, not "".
            # str.join raises TypeError on a None element, so filter here.
            if text:
                page_texts.append(text)

    return "\n\n".join(page_texts)


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes, including table contents.

    Tables are walked in document order alongside paragraphs rather than
    appended afterwards. Order carries meaning here: the extraction prompt
    infers a skill's `source` from its surroundings, so a skills table hoisted
    out of position reads as a bare keyword list even when it sat under a job
    heading.

    Raises:
        CorruptDocumentError: If the bytes cannot be parsed as a DOCX.
    """
    buffer = BytesIO(file_bytes)

    # Two failure modes, both from this one call:
    #   BadZipFile -- the bytes are not a zip at all (a DOCX is a zip).
    #   KeyError   -- a valid zip missing the parts a Word document needs,
    #                 which python-docx surfaces as a failed lookup.
    # Catching a builtin as broad as KeyError is only safe because the try
    # covers exactly one call; widen it and this starts swallowing our bugs.
    try:
        document = docx.Document(buffer)
    except (BadZipFile, KeyError) as exc:
        raise CorruptDocumentError(
            "This file could not be read as a Word document. It may be "
            "corrupt, incomplete, or not actually a .docx."
        ) from exc

    return "\n\n".join(_extract_blocks(document))


def _iter_blocks(parent: object) -> Iterator[Paragraph | Table]:
    """Yield the paragraphs and tables of `parent` in document order.

    `document.paragraphs` and `document.tables` are each complete and each
    partial: neither reports where its members sat relative to the other's, so
    reading both and concatenating produces every word in the wrong order.
    Walking the underlying XML is the only way to recover the sequence.

    Cell paragraphs are deliberately not double-counted -- `document.paragraphs`
    returns only direct children of the body, so a paragraph inside a table is
    reached through the table and nowhere else.
    """
    element = parent._tc if isinstance(parent, _Cell) else parent.element.body

    for child in element.iterchildren():
        # Matching on the qualified tag rather than importing CT_P and CT_Tbl,
        # which are python-docx internals and have moved between releases.
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _extract_blocks(parent: object) -> list[str]:
    """Flatten one container's blocks into text, recursing through tables."""
    blocks: list[str] = []

    for block in _iter_blocks(parent):
        if isinstance(block, Paragraph):
            # Same truthiness test the paragraph-only version used: Word
            # documents are full of genuinely empty paragraphs used as
            # spacing, while a whitespace-only one is usually deliberate.
            if block.text:
                blocks.append(block.text)
        else:
            blocks.extend(_extract_table(block))

    return blocks


def _extract_table(table: Table) -> list[str]:
    """Render one table as a list of row strings.

    Rows whose cells are all single-line are joined with CELL_SEPARATOR, which
    keeps a two-column skills table readable as pairs. Rows containing a
    multi-line cell are emitted cell-per-block instead -- some resumes lay the
    entire document out in one invisible table, and forcing that onto one line
    per row would produce a wall of text with separators scattered through it.
    """
    rows: list[str] = []

    for row in table.rows:
        cells: list[str] = []
        # A horizontally merged cell is repeated by `row.cells` once for every
        # grid column it spans, all backed by the same element. Without this
        # check a merged section header is emitted two or three times.
        seen: set[int] = set()

        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))

            text = "\n\n".join(_extract_blocks(cell))
            if text:
                cells.append(text)

        if not cells:
            continue

        if any("\n" in cell for cell in cells):
            rows.extend(cells)
        else:
            rows.append(CELL_SEPARATOR.join(cells))

    return rows
