"""Shared test fixtures, and the isolation that has to happen before them.

Every test document is constructed in memory at test time rather than
committed as a binary fixture. Two reasons -- .gitignore deliberately keeps
resumes out of this repo because they contain personal data, and a generated
document makes its own defect obvious. A committed `corrupt.pdf` is opaque;
`truncated_pdf()` says exactly what is wrong with it.

**The first thing this module does is repoint the suite at its own database
and its own Redis keyspace**, and that has to happen at import rather than in
a fixture. `app.config.Settings` reads its values once, when it is imported,
and `app.db`'s engine and `app.queues`' connection pool are built from it --
so by the time any fixture runs it is already too late. pytest imports
conftest before it collects test modules, which is the only window there is.
"""

import os
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.parse import urlsplit, urlunsplit
from uuid import uuid4

import docx
import psycopg
import pytest
from dotenv import dotenv_values
from PIL import Image

# --- test isolation ---------------------------------------------------------
#
# Until this existed the suite ran against the development database and Redis
# db 0 -- the same ones `docker compose up` leaves four workers and a scheduler
# using. The result was a suite that failed a different handful of tests on
# every run while each one passed in isolation. Observed directly, from
# worker-bulk during a test run:
#
#     ingest: app.workers.tasks.process_job_url(11782)
#     warning  process_job_url: job 11782 is gone
#
# A live container had consumed a job a test enqueued, then found the row
# deleted by that test's teardown. The same sharing let any worker bump the
# `gate:hits` counters `TestContentHashGate` asserts on, and left
# `test_recommendations` recalling over the real 1,500-posting catalog.

_BACKEND = Path(__file__).resolve().parent.parent
_DEV_ENV = dotenv_values(_BACKEND / ".env")


def _dev_value(key: str) -> str:
    """The developer's real setting, from the environment or backend/.env."""
    return os.environ.get(key) or _DEV_ENV.get(key) or ""


def _with_database(url: str, suffix: str = "_test") -> str:
    parts = urlsplit(url)
    name = parts.path.lstrip("/")
    if name.endswith(suffix):
        return url
    return urlunsplit(parts._replace(path=f"/{name}{suffix}"))


def _with_redis_db(url: str, index: str = "1") -> str:
    return urlunsplit(urlsplit(url)._replace(path=f"/{index}"))


# Re-entrancy matters, and the failure without it is baffling rather than
# loud-in-a-useful-way. `from conftest import make_pdf` resolves to the module
# pytest already loaded, but `from tests.conftest import ...` builds a *second*
# module object and runs this file again -- at which point DATABASE_URL is
# already the test URL, the derivation produces itself, and the guard below
# fires claiming the test database is the development one. The sentinel makes
# the second pass a no-op instead.
_SENTINEL = "FITCHECK_TEST_ENV_ACTIVE"

if os.environ.get(_SENTINEL) == "1":
    TEST_DATABASE_URL = os.environ["DATABASE_URL"]
    TEST_REDIS_URL = os.environ["REDIS_URL"]
else:
    _DEV_DATABASE_URL = _dev_value("DATABASE_URL")
    _DEV_REDIS_URL = _dev_value("REDIS_URL")

    if not _DEV_DATABASE_URL or not _DEV_REDIS_URL:
        raise RuntimeError(
            "DATABASE_URL and REDIS_URL must be set (backend/.env) so the test "
            "database can be derived from them. Run pytest from backend/."
        )

    # Derived by default so a fresh clone needs no extra configuration, and
    # overridable for CI, where the test database is usually somewhere else
    # entirely rather than a sibling of a developer's.
    TEST_DATABASE_URL = (
        os.environ.get("TEST_DATABASE_URL") or _with_database(_DEV_DATABASE_URL)
    )
    TEST_REDIS_URL = os.environ.get("TEST_REDIS_URL") or _with_redis_db(_DEV_REDIS_URL)

    # The guard that makes the rest of this safe to be aggressive about. Every
    # destructive thing below -- CREATE DATABASE, FLUSHDB, `alembic upgrade` --
    # is only defensible because it cannot be pointed at the real one. A
    # misconfigured TEST_DATABASE_URL should stop the run, not silently
    # truncate somebody's catalog.
    if TEST_DATABASE_URL == _DEV_DATABASE_URL:
        raise RuntimeError(
            f"refusing to run: the test database resolves to the development "
            f"one ({_DEV_DATABASE_URL!r}). Set TEST_DATABASE_URL to something "
            f"else."
        )
    if TEST_REDIS_URL == _DEV_REDIS_URL:
        raise RuntimeError(
            f"refusing to run: the test Redis resolves to the development one "
            f"({_DEV_REDIS_URL!r}). Set TEST_REDIS_URL to a different db index."
        )
    if urlsplit(TEST_REDIS_URL).path in ("", "/", "/0"):
        raise RuntimeError(
            "refusing to run: the test Redis must not be db 0, which holds the "
            "development sessions and queues."
        )

    # Environment variables win over the .env file in pydantic-settings, so
    # this is what redirects app.config, app.db, app.queues, and
    # alembic/env.py -- all of which read `settings` -- without editing a file
    # the running API and workers also read.
    os.environ["DATABASE_URL"] = TEST_DATABASE_URL
    os.environ["REDIS_URL"] = TEST_REDIS_URL
    os.environ[_SENTINEL] = "1"


def _admin_dsn(url: str) -> str:
    """A libpq URI for the `postgres` maintenance database on the same server.

    `CREATE DATABASE` cannot run from inside the database being created, so
    this is the one connection that necessarily goes somewhere else. It is
    also why the scheme loses `+psycopg`: that suffix is SQLAlchemy's dialect
    selector and psycopg does not understand it.
    """
    parts = urlsplit(url)
    return urlunsplit(parts._replace(scheme="postgresql", path="/postgres"))


@pytest.fixture(scope="session", autouse=True)
def _isolated_backing_services() -> None:
    """Create and migrate the test database; clear the test Redis keyspace.

    Session-scoped and autouse, so no test can opt out and forget. `alembic
    upgrade head` rather than `Base.metadata.create_all` deliberately: the
    schema then comes from the same migrations production runs, including the
    two `CREATE EXTENSION` statements and the partial and HNSW indexes that
    metadata alone would not reproduce. A suite passing against a schema
    nobody deploys is worth very little.
    """
    name = urlsplit(TEST_DATABASE_URL).path.lstrip("/")

    with psycopg.connect(_admin_dsn(TEST_DATABASE_URL), autocommit=True) as conn:
        exists = conn.execute(
            "SELECT 1 FROM pg_database WHERE datname = %s", (name,)
        ).fetchone()
        if not exists:
            # Identifier, so it cannot be parameterised. Safe because the name
            # is derived from a URL this process just built, not from input.
            conn.execute(f'CREATE DATABASE "{name}"')

    from alembic import command
    from alembic.config import Config

    config = Config(str(_BACKEND / "alembic.ini"))
    # Absolute, so the run does not depend on the working directory.
    config.set_main_option("script_location", str(_BACKEND / "alembic"))
    command.upgrade(config, "head")

    # Truncated once per session, for the rows no fixture owns.
    #
    # Most cleanup here rides on deleting a user and letting the cascade take
    # their profiles, jobs, and matches. `job_postings` has no owner by design
    # -- it is the shared catalog both paths converge on -- so nothing removes
    # it, and it accumulated across runs. That is not harmless: recall orders
    # by vector distance across the *whole* table, so a suite that had been run
    # fifty times would be reranking against a growing pile of other tests'
    # fixtures. `sources` and `url_batches` drift the same way.
    #
    # RESTART IDENTITY so ids do not climb without bound, CASCADE because the
    # tables reference each other. alembic_version is deliberately spared --
    # dropping it would re-run every migration on the next session.
    from sqlalchemy import create_engine, text

    engine = create_engine(TEST_DATABASE_URL)
    with engine.begin() as connection:
        tables = [
            row[0]
            for row in connection.execute(
                text(
                    "SELECT tablename FROM pg_tables "
                    "WHERE schemaname = 'public' AND tablename <> 'alembic_version'"
                )
            )
        ]
        if tables:
            listed = ", ".join(f'"{name}"' for name in tables)
            connection.execute(text(f"TRUNCATE {listed} RESTART IDENTITY CASCADE"))
    engine.dispose()

    # Flushed once per session rather than never. Redis holds process-global
    # state that no fixture owns -- the `gate:hits` / `gate:misses` counters,
    # session keys, RQ registries -- and leftovers from a previous run are
    # exactly what made TestContentHashGate flaky. Safe only because of the
    # db-0 guard above.
    from redis import Redis

    connection = Redis.from_url(TEST_REDIS_URL)
    connection.flushdb()
    connection.close()


def _escape_pdf_text(text: str) -> bytes:
    """Escape the three characters that are syntax inside a PDF string."""
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    return escaped.encode("ascii")


def make_pdf(text: str = "Hello") -> bytes:
    """Build a minimal single-page PDF with a real text layer.

    Hand-assembled rather than pulled from a library because nothing in
    requirements.txt writes PDFs, and adding a dependency to produce five
    test objects is a poor trade. The byte offsets in the cross-reference
    table are computed as the file is built -- pdfminer validates that table
    on open(), so they have to be right or every test here would exercise the
    corrupt-document path by accident.
    """
    stream = b"BT /F1 24 Tf 72 700 Td (" + _escape_pdf_text(text) + b") Tj ET"

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length "
        + str(len(stream)).encode()
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(number).encode() + b" 0 obj\n" + body + b"\nendobj\n"

    xref_offset = len(out)
    size = len(objects) + 1

    out += b"xref\n0 " + str(size).encode() + b"\n"
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("ascii")

    out += b"trailer\n<< /Size " + str(size).encode() + b" /Root 1 0 R >>\n"
    out += b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF\n"

    return bytes(out)


def make_scanned_pdf() -> bytes:
    """Build a PDF whose only content is an image -- no text layer at all.

    This is what a photographed or scanned resume looks like to pdfplumber:
    a structurally valid PDF where every page's extract_text() returns None.
    Pillow is already a pdfplumber dependency, so this costs nothing.
    """
    buffer = BytesIO()
    Image.new("RGB", (200, 200), "white").save(buffer, format="PDF")
    return buffer.getvalue()


def make_docx(paragraphs: list[str]) -> bytes:
    """Build a real .docx containing the given paragraphs."""
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_table(
    rows: list[list[str]],
    before: list[str] | None = None,
    after: list[str] | None = None,
) -> bytes:
    """Build a .docx with a table, optionally surrounded by paragraphs.

    `before` and `after` exist to pin document *order*. A table appended to the
    end of the file cannot distinguish an implementation that walks the body in
    sequence from one that reads paragraphs and tables separately and
    concatenates -- both pass. Content on either side of the table is what
    makes those two behaviours produce different output.
    """
    document = docx.Document()

    for text in before or []:
        document.add_paragraph(text)

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for cell_index, text in enumerate(row):
            # Cells start life holding one empty paragraph, so assigning .text
            # replaces it rather than appending a second.
            table.cell(row_index, cell_index).text = text

    for text in after or []:
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_merged_row(header: str, body: list[str]) -> bytes:
    """Build a table whose first row is one cell merged across every column.

    The layout a resume uses for a section banner above its columns. python-docx
    reports a merged cell once per grid column it spans, so this is the fixture
    that catches an implementation emitting the header N times.
    """
    document = docx.Document()

    table = document.add_table(rows=2, cols=len(body))
    merged = table.cell(0, 0)
    for column in range(1, len(body)):
        merged = merged.merge(table.cell(0, column))
    merged.text = header

    for column, text in enumerate(body):
        table.cell(1, column).text = text

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_nested_table(outer: str, inner: list[str]) -> bytes:
    """Build a table containing another table -- a cell holding a sub-grid.

    Rare in hand-written resumes and routine in ones exported from a template.
    Either the extractor recurses or the inner content vanishes silently.
    """
    document = docx.Document()

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = outer

    host = table.cell(0, 1)
    nested = host.add_table(rows=1, cols=len(inner))
    for column, text in enumerate(inner):
        nested.cell(0, column).text = text

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_zip_without_word_parts() -> bytes:
    """A structurally valid zip that is not a Word document.

    A .docx *is* a zip, so this passes the BadZipFile check and fails later,
    when python-docx looks for the document parts that are not there. It is
    the second of the two failure modes _extract_docx catches.
    """
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("not-a-word-document.txt", "hello")
    return buffer.getvalue()


@pytest.fixture
def make_user():
    """Create real user rows, cleaned up afterwards.

    Deleting the user cascades to their profiles and, through those, to their
    jobs -- so tests that create data under a user need no cleanup of their
    own.

    Emails are randomised because the unique index is case-insensitive and
    real: a fixed address would make the second test in a session collide
    with leftovers from the first.

    The domain is .dev rather than .local or .test on purpose -- those are
    IANA special-use domains and EmailStr rejects them, so a fixture using
    one would create accounts the API itself would refuse to register.
    """
    from app.db import SessionLocal
    from app.models import User
    from app.security import hash_password

    created: list[int] = []

    def _make(email: str | None = None, password: str = "testpassword123"):
        db = SessionLocal()
        try:
            user = User(
                email=email or f"t{uuid4().hex[:12]}@fitcheck.dev",
                password_hash=hash_password(password),
            )
            db.add(user)
            db.commit()
            db.refresh(user)
            created.append(user.id)
            db.expunge(user)
            return user
        finally:
            db.close()

    yield _make

    db = SessionLocal()
    try:
        for user_id in created:
            obj = db.get(User, user_id)
            if obj is not None:
                db.delete(obj)
        db.commit()
    finally:
        db.close()


@pytest.fixture
def as_user():
    """Authenticate the test client as a given user.

    Overrides the `current_user` dependency rather than performing a real
    login, so tests that are about something else do not each need Redis and
    a cookie jar. `owned_profile` is deliberately *not* overridden -- it
    depends on `current_user`, so the real ownership predicate still runs and
    is still what the authorization tests exercise.

    The genuine cookie flow is covered end to end in test_auth.py.
    """
    from app.main import app
    from app.security import current_user

    def _as(user):
        app.dependency_overrides[current_user] = lambda: user
        return user

    yield _as

    app.dependency_overrides.clear()


@pytest.fixture
def anyio_backend() -> str:
    """Run @pytest.mark.anyio tests on asyncio only.

    anyio's pytest plugin parameterises over asyncio and trio by default;
    trio is not installed and this app only ever runs on asyncio.
    """
    return "asyncio"
